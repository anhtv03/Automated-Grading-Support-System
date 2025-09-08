import os
import re
from typing import List, Type, Callable

from docx import Document
from fastapi import UploadFile, File, Depends
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session

from app.constants.file_type import FileType
from app.constants.status import FileUploadType, SemesterStatus
from app.db import db_answer_template
from app.db.database import get_db
from app.db.db_answer_template import get_all_answer_template_by_session_id
from app.db.db_exam import db_update_exam
from app.db.db_grading_guide import get_grading_guide_by_session_id
from app.db.db_submission import get_all_submissions_by_session_id
from app.db.db_submission_question import (create_submission_questions)
from app.db.models import AnswerTemplate, SubmissionQuestion
from app.exceptions.custom_exception import CustomException
from app.exceptions.error_codes import ErrorCode
from app.external.aws_service import S3Uploader
from app.schemas.sche_base import BaseRequest, RequestT, ResponseT
from app.schemas.sche_exam_template import AnswerTemplateResponse
from app.schemas.sche_submission import SubmissionCreate
from app.schemas.sche_submisson_question import SubmissionQuestionCreateRequest, SubmissionQuestionResponse
from app.services.answer_template_service import AnswerTemplateService
from app.services.exam_question_service import ExamQuestionService
from app.services.grading_guide_service import GradingGuideService
from app.services.submission_question_service import QuestionService
from app.services.submission_service import SubmissionService
from app.services.upload_session_service import UploadSessionService
from app.utils.file_util import is_word_or_excel
from app.utils.text_util import check_content_submission
from app.utils.word_util import read_file_word, read_file_excel, get_number_of_question_by_content


class FileService:
    QUESTION_PATTERNS = [
        r"(?i)^question\s*\d+",
        r"(?i)^q\s*\d+"
    ]

    @staticmethod
    def extract_ordered_blocks(docx_path):
        try:
            doc = Document(docx_path)
            blocks = []

            def add_table(table):
                rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        rows.append(" | ".join(row_data))
                if rows:
                    blocks.append("\n".join(rows))

            for block in doc.element.body:
                if block.tag.endswith("p"):
                    para = next((p for p in doc.paragraphs if p._p == block), None)
                    if para and para.text.strip():
                        blocks.append(para.text.strip())
                elif block.tag.endswith("tbl"):
                    table = next((t for t in doc.tables if t._tbl == block), None)
                    if table:
                        add_table(table)

            return blocks
        except CustomException as e:
            raise
        except Exception as e:
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

    @staticmethod
    def match_question_header(text):
        try:
            for pattern in FileService.QUESTION_PATTERNS:
                if re.match(pattern, text.strip()):
                    return True
            return False
        except CustomException as e:
            raise
        except Exception as e:
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

    @staticmethod
    def extract_question_name(full_question):
        try:
            match = re.search(r"(?i)(question|q)?\s*(\d+)", full_question)
            if match:
                number = match.group(2)
                return f"Question {number}"
            return full_question.strip()
        except CustomException as e:
            raise
        except Exception as e:
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

    @staticmethod
    def group_blocks_by_question(blocks):
        # try:
        grouped = []
        current_question = None
        current_content = []

        for block in blocks:
            if FileService.match_question_header(block):
                if current_question:
                    grouped.append({
                        "questionName": FileService.extract_question_name(current_question),
                        "content": "\n".join(current_content).strip()
                    })
                current_question = block.strip()
                current_content = []
            elif current_question:
                current_content.append(block)

        if current_question:
            grouped.append({
                "questionName": FileService.extract_question_name(current_question),
                "content": "\n".join(current_content).strip()
            })

        return grouped

    # except CustomException as e:
    #     raise
    # except Exception as e:
    #     raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

    @staticmethod
    async def create_document_files(
            request: BaseRequest,
            files: List[UploadFile] = File(...),
            type: str = FileUploadType.ANSWER_TEMPLATE,
            request_class: Type[RequestT] = AnswerTemplate,
            response_class: Type[ResponseT] = AnswerTemplateResponse,
            db_create_func: Callable[[Session, RequestT], RequestT] = db_answer_template.create_answer_template,
            db: Session = Depends(get_db)
    ) -> List[ResponseT]:
        try:
            names = request.name.split(",") if request.name else []
            index = 0
            results = []

            for file in files:
                try:
                    # Validate and read file
                    is_valid_file = await is_word_or_excel(file)
                    content = ""

                    if is_valid_file == FileType.Word:
                        content = await read_file_word(file)
                    elif is_valid_file == FileType.Excel:
                        content = await read_file_excel(file)

                    clear_content = check_content_submission(content)
                    print(clear_content)

                    # Upload to S3
                    file_key = ""
                    if content and content.strip():
                        result = await S3Uploader.upload_file_to_s3(file)
                        file_key = result['file_key']

                    # Create object
                    new_object = request_class(
                        name=names[index],
                        session_id=request.session_id,
                        file_key=file_key,
                        content=content,
                    )

                    match type:
                        case FileUploadType.ANSWER_TEMPLATE:
                            number_question = get_number_of_question_by_content(content)
                            guide = get_grading_guide_by_session_id(db, request.session_id)
                            if guide:
                                if guide.question_number != number_question:
                                    raise CustomException(ErrorCode.NUMBER_QUESTION_NOT_MATCH)
                                elif guide.question_number == 0 & number_question == 0:
                                    raise CustomException(ErrorCode.NUMBER_QUESTION_NOT_MATCH)

                            new_object.question_number = number_question

                        case FileUploadType.GUIDE_TEMPLATE:
                            number_question = get_number_of_question_by_content(content)
                            answer = get_all_answer_template_by_session_id(db, request.session_id)
                            if answer:
                                if answer.question_number != number_question:
                                    raise CustomException(ErrorCode.NUMBER_QUESTION_NOT_MATCH)
                                elif answer.question_number == 0 & number_question == 0:
                                    raise CustomException(ErrorCode.NUMBER_QUESTION_NOT_MATCH)

                            new_object.question_number = number_question

                    # Save to database using provided function
                    created_object = await db_create_func(db, new_object)

                    if type == FileUploadType.EXAM:
                        ExamQuestionService.create_exam_question_by_exam_content(created_object.id, content, db)
                        created_object.status = SemesterStatus.VISIBLE
                        db_update_exam(db, created_object)
                    elif type == FileUploadType.GUIDE_TEMPLATE:
                        criteria = GradingGuideService.extract_grading_guide_criteria(db, content, created_object.id)
                        print(criteria)

                    # Convert to response class
                    results.append(response_class.model_validate(created_object))

                except CustomException as e:
                    # if e.message == ErrorCode.NUMBER_QUESTION_NOT_MATCH.message:
                    raise
                # else:
                #     continue
                except Exception as e:
                    print(f"Error processing file {file.filename}: {e}")
                    continue
                index += 1

            return results
        except CustomException as e:
            raise
        except Exception as e:
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

    @staticmethod
    async def process_docx_upload(file: UploadFile, temp_prefix: str = "temp") -> dict:
        """
        - Ghi file tạm
        - Tách nội dung từ file docx
        - Upload lên S3
        - Trả về content và file_key (đường dẫn trên S3)
        """
        try:
            temp_filename = f"{file.filename}"
            temp_path = os.path.join(temp_prefix, temp_filename)

            content_bytes = await file.read()

            with open(temp_path, 'wb') as f:
                f.write(content_bytes)

            blocks = FileService.extract_ordered_blocks(temp_path)
            content = "\n".join(blocks).strip()

            file_key = ""
            if content:
                result = await S3Uploader.upload_file_to_s3(file)
                file_key = result["file_key"]

            return {
                "blocks": blocks,
                "content": content,
                "file_key": file_key,
                "temp_path": temp_path
            }

        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    @staticmethod
    async def create_submission(
            session_id: int,
            submission_names: List[str],
            submission_files: List[UploadFile],
            user_id: int,
            db: Session
    ):
        if len(submission_names) != len(submission_files):
            raise ValueError("Số lượng tên và file submission không khớp")

        try:
            # Bắt đầu transaction
            with db.begin():
                # Lấy session upload của user
                upload_session = UploadSessionService.get_upload_session_by_session_id(db, session_id, user_id)

                question_names = []
                template_exam = AnswerTemplateService.get_exam_template_by_session_id(session_id, db)

                if template_exam:
                    blocks = template_exam.content.split("\n")
                    question_names = [q["questionName"] for q in FileService.group_blocks_by_question(blocks)]

                # ===== Xử lý từng Submission =====
                for name, file in zip(submission_names, submission_files):
                    result = await FileService.process_docx_upload(file)

                    subCreate = SubmissionCreate(
                        session_id=session_id,
                        name=name,
                        content=result["content"],
                        file_key=result["file_key"],
                    )
                    submission = await SubmissionService.save_submission(db, subCreate)

                    for question_name in question_names:
                        new_submission_question = SubmissionQuestionCreateRequest(
                            submission_id=submission.id,
                            question_name=question_name
                        )
                        await QuestionService.save_submission_question(db, new_submission_question)

        except SQLAlchemyError as e:
            db.rollback()
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

        except Exception as e:
            db.rollback()
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

    @staticmethod
    def create_submission_questions(db: Session, session_id: int) -> List[SubmissionQuestionResponse]:
        # try:
        question_names = []
        template_exam = AnswerTemplateService.get_exam_template_by_session_id(session_id, db)

        if template_exam:
            blocks = template_exam.content.split("\n")
            question_names = [q["questionName"] for q in FileService.group_blocks_by_question(blocks)]

        submissions = get_all_submissions_by_session_id(db, session_id)
        submission_questions = []
        # if no question name so no extract insert
        for submission in submissions:
            for question_name in question_names:
                new_submission_question = SubmissionQuestion(
                    submission_id=submission.id,
                    question_name=question_name
                )
                submission_questions.append(new_submission_question)

        data = create_submission_questions(db, submission_questions)

        return [SubmissionQuestionResponse.model_validate(q.__dict__) for q in data]

    # except CustomException as e:
    #     raise
    # except Exception as e:
    #     raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)
    @staticmethod
    async def get_document_files_by_session_id(
            session_id: int,
            response_class: Type[ResponseT] = AnswerTemplateResponse,
            db_get_func: Callable[[Session, RequestT], RequestT] = db_answer_template.create_answer_template,
            db: Session = Depends(get_db)
    ) -> List[ResponseT]:
        try:
            results = []
            search_data = db_get_func(db, session_id)

            # Convert to response class
            for item in search_data:
                results.append(response_class.model_validate(item))

            return results
        except CustomException as e:
            raise
        except Exception as e:
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)

    @staticmethod
    async def delete_document_files_by_id(
            ids: List[int],
            db_delete_func: Callable[[Session, RequestT], RequestT] = db_answer_template.db_delete_exam_templates,
            db: Session = Depends(get_db)
    ) -> int:
        try:
            count = await db_delete_func(db, ids)

            return count
        except CustomException as e:
            raise
        except Exception as e:
            raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)
