import io
import os
import re

import pypandoc
from docx import Document
from fastapi import UploadFile

from app.exceptions.custom_exception import CustomException
from app.exceptions.error_codes import ErrorCode


def validate_docx(file: UploadFile, content: bytes) -> None:
    if not (file.filename.lower().endswith('.docx') or file.filename.lower().endswith('.doc')):
        raise CustomException(ErrorCode.FILE_NOT_DOCX)

    if len(content) > 5 * 1024 * 1024:
        raise CustomException(ErrorCode.FILE_SIZE_EXCEEDED)

    doc = Document(io.BytesIO(content))
    if not doc.paragraphs and not doc.tables:
        raise CustomException(ErrorCode.FILE_EMPTY)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            raise CustomException(ErrorCode.FILE_CONTAINS_IMAGES)

    for shape in doc.inline_shapes:
        raise CustomException(ErrorCode.FILE_CONTAINS_SHAPES)

    for para in doc.paragraphs:
        try:
            para.text.encode('utf-8')
        except UnicodeEncodeError:
            raise CustomException(ErrorCode.FILE_INVALID_CHARACTERS)


async def read_file_word(file: UploadFile, temp_prefix: str = "temp") -> str:
    """
    Read the content of a .docx file, validate its format, and return the text content.
    """
    try:
        # Ensure temp folder exists
        os.makedirs(temp_prefix, exist_ok=True)

        # Create and save temporary file
        temp_filename = f"{file.filename}"
        temp_path = os.path.join(temp_prefix, temp_filename)
        content_bytes = await file.read()

        if file.filename.lower().endswith('.doc'):
            temp_doc = 'temp.doc'
            temp_docx = 'temp.docx'
            with open(temp_doc, 'wb') as f:
                f.write(content_bytes)

            pypandoc.convert_file(temp_doc, 'docx', outputfile=temp_docx)
            with open(temp_docx, 'rb') as f:
                content_bytes = f.read()
            if os.path.exists(temp_doc):
                os.remove(temp_doc)
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

        with open(temp_path, 'wb') as f:
            f.write(content_bytes)

        # Validate file format
        validate_docx(file, content_bytes)

        # Check file content
        try:
            doc = Document(temp_path)
            if not doc.core_properties.title and not doc.paragraphs:
                raise CustomException(ErrorCode.FILE_INVALID_FORMAT)
        except Exception:
            raise CustomException(ErrorCode.FILE_INVALID_FORMAT)

        # Extract content
        blocks = extract_ordered_blocks(temp_path)
        content = "\n".join(blocks).strip()

        return content

    finally:
        # Remove temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)


async def read_file_excel(file: UploadFile) -> str:
    # Giả định hàm này sẽ được định nghĩa để đọc file Excel
    # Bạn có thể sử dụng pandas hoặc openpyxl để đọc file Excel
    return "sheet"
    # raise NotImplementedError("read_file_excel function is not implemented yet.")

def extract_ordered_blocks(docx_path):
    try:
        doc = Document(docx_path)
        blocks = []

        def add_table(table):
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    rows.append(" | ".join(row_data))
            if rows:
                blocks.append("\n".join(rows))

        for block in doc.element.body:
            if block.tag.endswith("p"):
                para = next((p for p in doc.paragraphs if p._p == block), None)
                if para and para.text.strip():
                    blocks.append(para.text.strip())
            elif block.tag.endswith("tbl"):
                table = next((t for t in doc.tables if t._tbl == block), None)
                if table:
                    add_table(table)

        return blocks
    except CustomException as e:
        raise
    except Exception as e:
        raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)


@staticmethod
def get_number_of_question_by_content(
        content: str
) -> int:
    try:
        matches = re.findall(r'(?i)^Question\s*\d+', content, re.MULTILINE)

        return len(matches)
    except CustomException as e:
        raise
    except Exception as e:
        raise CustomException(ErrorCode.INTERNAL_SERVER_ERROR)
